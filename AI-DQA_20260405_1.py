import streamlit as st
import pandas as pd
import json
import os
import sqlite3
import openai
import re
import secrets
import string
import stripe
from io import BytesIO
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Tuple
from duckduckgo_search import DDGS
from neo4j import GraphDatabase
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ================== 页面配置 ==================
st.set_page_config(page_title="AI+DQA 风险分析系统", page_icon="🔍", layout="wide")

# ================== 试用模式的安全防护代码（CSS + JS）水印加大版 ==================
TRIAL_SECURITY_HTML = """
<style>
    body, .stApp, .report-card, .markdown-text-container {
        user-select: none !important;
        -webkit-user-select: none !important;
        -moz-user-select: none !important;
        -ms-user-select: none !important;
        -webkit-touch-callout: none !important;
    }
    .trial-watermark-bg {
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        pointer-events: none;
        z-index: 9999;
        background-image: repeating-linear-gradient(45deg, 
            rgba(0,0,0,0.05) 0px, rgba(0,0,0,0.05) 4px,
            transparent 4px, transparent 60px,
            rgba(0,0,0,0.05) 60px, rgba(0,0,0,0.05) 64px,
            transparent 64px, transparent 120px);
        background-size: 120px 120px;
    }
    .trial-watermark-text {
        position: fixed;
        bottom: 20px;
        right: 20px;
        opacity: 0.5;
        font-size: 14px;
        color: #666;
        background: rgba(255,255,255,0.8);
        padding: 8px 16px;
        border-radius: 8px;
        font-family: monospace;
        pointer-events: none;
        z-index: 10000;
        width: 360px;
        max-width: 80%;
        text-align: right;
        box-shadow: 0 2px 6px rgba(0,0,0,0.1);
    }
</style>
<script>
    document.addEventListener('contextmenu', function(e) { e.preventDefault(); return false; });
    document.addEventListener('keydown', function(e) {
        if (e.ctrlKey && (e.key === 'c' || e.key === 'C' || e.key === 'v' || e.key === 'V' || 
                          e.key === 'x' || e.key === 'X' || e.key === 's' || e.key === 'S')) {
            e.preventDefault(); return false;
        }
        if (e.key === 'F12') { e.preventDefault(); return false; }
    });
    document.addEventListener('selectstart', function(e) { e.preventDefault(); return false; });
</script>
<div class="trial-watermark-bg"></div>
<div class="trial-watermark-text">⚠️ 机密报告 · 请联系 Techlife2027@gmail.com 购买授权 ⚠️</div>
"""

# ================== Stripe 配置 ==================
try:
    stripe.api_key = st.secrets["STRIPE_SECRET_KEY"]
except:
    stripe.api_key = ""

# 套餐定义（次数，有效期月数）—— 已填入用户提供的 Price ID
PLANS = {
    "single": {"uses": 3, "months": 9999, "price_id": "price_1TLmO14PvqyeiHq5I1JiQUBe", "name_zh": "单次通行", "name_en": "Single Pass", "price_usd": 3},
    "50": {"uses": 50, "months": 1, "price_id": "price_1TLgfP4PvqyeiHq5etIezp0y", "name_zh": "50次套餐", "name_en": "50 Credits", "price_usd": 30},
    "1000": {"uses": 1000, "months": 12, "price_id": "price_1TLgfQ4PvqyeiHq5FzEr7r71", "name_zh": "1000次套餐", "name_en": "1000 Credits", "price_usd": 200},
}

# ================== 授权与试用数据管理 ==================
USAGE_FILE = "usage_data.json"

def load_usage_data():
    if os.path.exists(USAGE_FILE):
        try:
            with open(USAGE_FILE, "r") as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_usage_data(data):
    with open(USAGE_FILE, "w") as f:
        json.dump(data, f, indent=2)

LICENSE_TYPES = {
    "trial": {"name": "试用版", "max_uses": 3, "max_months": 1, "en_name": "Trial"},
    "level1": {"name": "一级用户", "max_uses": 100, "max_months": 12, "en_name": "Level 1"},
    "level2": {"name": "二级用户", "max_uses": 300, "max_months": 24, "en_name": "Level 2"},
    "level3": {"name": "三级用户", "max_uses": 500, "max_months": 36, "en_name": "Level 3"},
    "level4": {"name": "四级用户", "max_uses": 1000, "max_months": 60, "en_name": "Level 4"},
}

def generate_report_key(license_type, custom_uses=None, custom_months=None, custom_key=None):
    if license_type == "custom":
        max_uses = custom_uses
        max_months = custom_months
        type_name = "自定义"
    else:
        lic_info = LICENSE_TYPES[license_type]
        max_uses = lic_info["max_uses"]
        max_months = lic_info["max_months"]
        type_name = lic_info["name"]
    expiry = datetime.now() + timedelta(days=max_months*30)
    expiry_str = expiry.isoformat()
    usage_db = load_usage_data()
    if custom_key and custom_key.strip():
        new_key = custom_key.strip().upper()
        if new_key in usage_db:
            return None, 0, None, "授权码已存在"
    else:
        while True:
            random_str = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for _ in range(8))
            new_key = f"{license_type.upper()}_{random_str}"
            if new_key not in usage_db:
                break
    usage_db[new_key] = {
        "type": license_type,
        "remaining": max_uses,
        "expiry": expiry_str,
        "total_uses": 0,
        "generated_at": datetime.now().isoformat()
    }
    save_usage_data(usage_db)
    return new_key, max_uses, expiry_str, type_name

def activate_license(report_key):
    if not report_key:
        return False, 0, None, None
    usage_db = load_usage_data()
    if report_key in usage_db:
        record = usage_db[report_key]
        remaining = record["remaining"]
        expiry_str = record["expiry"]
        expiry = datetime.fromisoformat(expiry_str)
        if remaining > 0 and datetime.now() <= expiry:
            return True, remaining, expiry_str, record.get("type", "unknown")
    return False, 0, None, None

def consume_usage(report_key):
    if st.session_state.get("admin_logged_in", False):
        return True
    if not report_key:
        if st.session_state.trial_uses_left > 0:
            st.session_state.trial_uses_left -= 1
            return True
        else:
            return False
    usage_db = load_usage_data()
    if report_key in usage_db:
        record = usage_db[report_key]
        if record["remaining"] > 0 and datetime.now() <= datetime.fromisoformat(record["expiry"]):
            record["remaining"] -= 1
            record["total_uses"] = record.get("total_uses", 0) + 1
            save_usage_data(usage_db)
            return True
    return False

def get_remaining_info(report_key):
    if st.session_state.get("admin_logged_in", False):
        return ("无限" if st.session_state.lang=="zh" else "Unlimited"), ("永久" if st.session_state.lang=="zh" else "Permanent")
    if report_key:
        valid, remaining, expiry_str, _ = activate_license(report_key)
        if valid:
            return str(remaining), expiry_str[:10]
    return str(st.session_state.trial_uses_left), ("试用剩余次数" if st.session_state.lang=="zh" else "Trial left")

def is_premium_user(report_key):
    if st.session_state.get("admin_logged_in", False):
        return True
    if report_key:
        valid, _, _, _ = activate_license(report_key)
        return valid
    return False

# ================== 初始化 Session State ==================
if "lang" not in st.session_state:
    st.session_state.lang = "zh"
if "admin_logged_in" not in st.session_state:
    st.session_state.admin_logged_in = False
if "enable_web_search" not in st.session_state:
    st.session_state.enable_web_search = True
if "translation_cache" not in st.session_state:
    st.session_state.translation_cache = {}
if "temp_api_key" not in st.session_state:
    st.session_state.temp_api_key = ""
if "temp_base_url" not in st.session_state:
    st.session_state.temp_base_url = "https://api.deepseek.com"
if "temp_model" not in st.session_state:
    st.session_state.temp_model = "deepseek-chat"
if "analyst_name" not in st.session_state:
    st.session_state.analyst_name = ""
if "analyst_title" not in st.session_state:
    st.session_state.analyst_title = ""
if "current_report_key" not in st.session_state:
    st.session_state.current_report_key = ""
if "trial_uses_left" not in st.session_state:
    st.session_state.trial_uses_left = 3
if "report_content" not in st.session_state:
    st.session_state.report_content = None
if "last_product_name" not in st.session_state:
    st.session_state.last_product_name = ""
if "last_product_desc" not in st.session_state:
    st.session_state.last_product_desc = ""
if "show_payment_dialog" not in st.session_state:
    st.session_state.show_payment_dialog = False
if "payment_new_key" not in st.session_state:
    st.session_state.payment_new_key = ""

ADMIN_USERNAME = "Laurence_ku"
ADMIN_PASSWORD = "Ku_product$2026"

# ================== 数据库部分（完整实现） ==================
class RiskDatabase:
    def get_risks(self, product_type: str) -> List[Dict]:
        raise NotImplementedError
    def get_product_decomposition(self, product_name: str, description: str) -> Dict:
        raise NotImplementedError
    def get_mitigation(self, module: str, failure_mode: str) -> str:
        raise NotImplementedError
    def get_knowledge_by_category(self, category: str) -> List[str]:
        raise NotImplementedError
    def add_knowledge(self, category: str, content: str) -> None:
        raise NotImplementedError
    def delete_knowledge(self, category: str, content: str) -> None:
        raise NotImplementedError
    def clear_knowledge_category(self, category: str) -> None:
        raise NotImplementedError
    def get_all_knowledge(self) -> Dict[str, List[str]]:
        raise NotImplementedError
    def load_initial_data(self) -> None:
        raise NotImplementedError
    def search_knowledge(self, keywords: str, limit: int = 5) -> List[str]:
        raise NotImplementedError

class SQLiteDatabase(RiskDatabase):
    def __init__(self):
        self.conn = sqlite3.connect('app_data.db', check_same_thread=False)
        self.init_tables()
        self.migrate_existing_knowledge()
        self.load_caches()

    def init_tables(self):
        cursor = self.conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS knowledge_base
                          (category TEXT, content TEXT, content_en TEXT)''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS product_risks
                          (product_type TEXT, module TEXT, failure_mode TEXT, cause TEXT,
                           severity INTEGER, occurrence INTEGER, detection INTEGER, mitigation TEXT)''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS industry_risks
                          (category TEXT, product_type TEXT, failure_mode TEXT, cause TEXT,
                           mitigation TEXT, source TEXT)''')
        cursor.execute("PRAGMA table_info(knowledge_base)")
        cols = [col[1] for col in cursor.fetchall()]
        if 'content_en' not in cols:
            cursor.execute("ALTER TABLE knowledge_base ADD COLUMN content_en TEXT")
        self.conn.commit()

    def migrate_existing_knowledge(self):
        cursor = self.conn.cursor()
        cursor.execute("SELECT rowid, category, content FROM knowledge_base WHERE content_en IS NULL OR content_en = ''")
        rows = cursor.fetchall()
        if not rows:
            return
        for rowid, cat, zh_text in rows:
            en_text = translate_text(zh_text, "en")
            cursor.execute("UPDATE knowledge_base SET content_en = ? WHERE rowid = ?", (en_text, rowid))
        self.conn.commit()

    def load_caches(self):
        cursor = self.conn.cursor()
        cursor.execute("SELECT category, content, content_en FROM knowledge_base")
        rows = cursor.fetchall()
        self.knowledge_zh = {"光学": [], "机械": [], "材料": [], "热学": [], "电气": [], "控制": []}
        self.knowledge_en = {"光学": [], "机械": [], "材料": [], "热学": [], "电气": [], "控制": []}
        for cat, zh, en in rows:
            if cat in self.knowledge_zh:
                self.knowledge_zh[cat].append(zh)
                if en:
                    self.knowledge_en[cat].append(en)
                else:
                    en_trans = translate_text(zh, "en")
                    self.knowledge_en[cat].append(en_trans)
                    cursor.execute("UPDATE knowledge_base SET content_en = ? WHERE category = ? AND content = ?", (en_trans, cat, zh))
        self.conn.commit()

        cursor.execute("SELECT product_type, module, failure_mode, cause, severity, occurrence, detection, mitigation FROM product_risks")
        rows = cursor.fetchall()
        self.product_risks = {}
        for row in rows:
            ptype = row[0]
            if ptype not in self.product_risks:
                self.product_risks[ptype] = []
            self.product_risks[ptype].append({
                "module": row[1], "failure_mode": row[2], "cause": row[3],
                "severity": row[4], "occurrence": row[5], "detection": row[6],
                "mitigation": row[7]
            })

    def get_risks(self, product_type: str) -> List[Dict]:
        risks = self.product_risks.get(product_type, [])
        for r in risks:
            r["RPN"] = r["severity"] * r["occurrence"] * r["detection"]
        return sorted(risks, key=lambda x: x["RPN"], reverse=True)[:10]

    def get_product_decomposition(self, product_name: str, description: str) -> Dict:
        if "路灯" in product_name or "street light" in product_name.lower():
            return {"product_type": "LED路灯", "function_units": ["光学","电气","热学"], "modules": ["LED光源","驱动电源"]}
        elif "天棚灯" in product_name or "high bay" in product_name.lower():
            return {"product_type": "高功率天棚灯", "function_units": ["光学","电气","热学","控制"], "modules": ["COB光源","风扇","热管"]}
        else:
            return {"product_type": "default", "function_units": ["电气","机械"], "modules": ["PCBA"]}

    def get_mitigation(self, module: str, failure_mode: str) -> str:
        keywords = f"{module} {failure_mode}"
        results = self.search_knowledge(keywords, limit=3)
        if results:
            return results[0][:200]
        return "建议参考行业规范和设计指南进行优化。"

    def get_knowledge_by_category(self, category: str) -> List[str]:
        lang = st.session_state.lang
        if lang == "zh":
            return self.knowledge_zh.get(category, [])
        else:
            return self.knowledge_en.get(category, [])

    def add_knowledge(self, category: str, content: str):
        lang = st.session_state.lang
        if lang == "zh":
            zh_text = content
            en_text = translate_text(content, "en")
        else:
            en_text = content
            zh_text = translate_text(content, "zh")
        cursor = self.conn.cursor()
        cursor.execute("INSERT INTO knowledge_base (category, content, content_en) VALUES (?, ?, ?)",
                       (category, zh_text, en_text))
        self.conn.commit()
        self.load_caches()

    def delete_knowledge(self, category: str, content: str):
        lang = st.session_state.lang
        if lang == "zh":
            cursor = self.conn.cursor()
            cursor.execute("DELETE FROM knowledge_base WHERE category = ? AND content = ?", (category, content))
        else:
            cursor = self.conn.cursor()
            cursor.execute("DELETE FROM knowledge_base WHERE category = ? AND content_en = ?", (category, content))
        self.conn.commit()
        self.load_caches()

    def clear_knowledge_category(self, category: str):
        cursor = self.conn.cursor()
        cursor.execute("DELETE FROM knowledge_base WHERE category = ?", (category,))
        self.conn.commit()
        self.load_caches()

    def get_all_knowledge(self) -> Dict[str, List[str]]:
        lang = st.session_state.lang
        if lang == "zh":
            return self.knowledge_zh
        else:
            return self.knowledge_en

    def search_knowledge(self, keywords: str, limit: int = 5) -> List[str]:
        if not keywords.strip():
            return []
        cursor = self.conn.cursor()
        query = """
            SELECT content, content_en FROM knowledge_base 
            WHERE content LIKE ? OR content_en LIKE ?
            LIMIT ?
        """
        like_pattern = f"%{keywords}%"
        cursor.execute(query, (like_pattern, like_pattern, limit))
        rows = cursor.fetchall()
        lang = st.session_state.lang
        results = []
        for row in rows:
            zh, en = row
            if lang == "zh":
                results.append(zh)
            else:
                results.append(en)
        return results

    def load_initial_data(self):
        cursor = self.conn.cursor()
        industry_data = [
            ("LED", "LED路灯", "光衰过快", "结温过高", "优化散热设计", "IEC 62031"),
            ("LED", "LED路灯", "浪涌损坏", "雷击", "加装SPD", "IEC 61643-11"),
            ("LED", "LED吸顶灯", "频闪", "纹波过大", "增加滤波", "IEEE 1789"),
            ("清洁电器", "洗地机", "滚刷堵转", "毛发缠绕", "防缠绕结构", "行业实践"),
            ("清洁电器", "吸尘器", "吸力下降", "滤网堵塞", "定期清理", "IEC 60312-1"),
            ("宠物电器", "宠物饮水机", "水泵噪音", "叶轮磨损", "无刷水泵", "行业标准"),
            ("宠物电器", "宠物喂食器", "卡粮", "粮食受潮", "干燥剂", "行业实践"),
        ]
        cursor.execute("DELETE FROM industry_risks")
        for row in industry_data:
            cursor.execute("INSERT INTO industry_risks (category, product_type, failure_mode, cause, mitigation, source) VALUES (?,?,?,?,?,?)", row)
        cursor.execute("SELECT COUNT(*) FROM product_risks")
        if cursor.fetchone()[0] == 0:
            default_risks = [
                ("LED路灯", "LED光源", "光衰过快", "结温过高", 8,7,5,"优化散热"),
                ("LED路灯", "驱动电源", "电容鼓包", "高温",9,6,6,"长寿命电容"),
                ("洗地机", "滚刷电机", "堵转", "毛发",8,7,6,"过流保护"),
                ("吸尘器", "电机", "吸力下降", "堵塞",7,6,5,"定期清理"),
                ("宠物饮水机", "水泵", "噪音", "磨损",6,5,4,"无刷电机"),
                ("宠物喂食器", "出粮机构", "卡粮", "受潮",8,5,6,"干燥剂"),
            ]
            for row in default_risks:
                cursor.execute("INSERT INTO product_risks VALUES (?,?,?,?,?,?,?,?)", row)
        self.conn.commit()
        self.load_caches()

class Neo4jDatabase(RiskDatabase):
    def __init__(self):
        self.driver = None
        self.connect()
        if self.driver:
            self._init_constraints()
            self._migrate_existing_knowledge()

    def connect(self):
        try:
            uri = st.secrets.get("NEO4J_URI", "")
            user = st.secrets.get("NEO4J_USERNAME", "neo4j")
            password = st.secrets.get("NEO4J_PASSWORD", "")
            if not uri or not password:
                return
            self.driver = GraphDatabase.driver(uri, auth=(user, password))
            with self.driver.session() as session:
                session.run("RETURN 1")
        except Exception:
            self.driver = None

    def _init_constraints(self):
        if not self.driver: return
        with self.driver.session() as session:
            try:
                session.run("CREATE CONSTRAINT knowledge_id IF NOT EXISTS FOR (k:Knowledge) REQUIRE k.id IS UNIQUE")
            except: pass
            try:
                session.run("CREATE INDEX knowledge_content IF NOT EXISTS FOR (k:Knowledge) ON (k.content)")
                session.run("CREATE INDEX knowledge_content_en IF NOT EXISTS FOR (k:Knowledge) ON (k.content_en)")
            except: pass

    def _migrate_existing_knowledge(self):
        if not self.driver: return
        with self.driver.session() as session:
            result = session.run("MATCH (k:Knowledge) WHERE k.content_en IS NULL RETURN k.category AS cat, k.content AS content, id(k) AS id")
            for rec in result:
                cat, zh_text, nid = rec["cat"], rec["content"], rec["id"]
                en_text = translate_text(zh_text, "en") if re.search(r'[\u4e00-\u9fff]', zh_text) else zh_text
                if not re.search(r'[\u4e00-\u9fff]', zh_text):
                    zh_text = translate_text(zh_text, "zh")
                session.run("MATCH (k:Knowledge) WHERE id(k)=$id SET k.content_en=$en, k.content=$zh",
                            {"id": nid, "en": en_text, "zh": zh_text})

    def _query(self, query, params=None):
        if not self.driver: return []
        with self.driver.session() as session:
            result = session.run(query, params or {})
            return [record.data() for record in result]

    def get_risks(self, product_type: str) -> List[Dict]:
        if not self.driver: return []
        cypher = """
            MATCH (p:ProductType {name: $ptype})-[:HAS_RISK]->(r:Risk)
            RETURN r.module AS module, r.failure_mode AS failure_mode, r.cause AS cause,
                   r.severity AS severity, r.occurrence AS occurrence, r.detection AS detection,
                   r.mitigation AS mitigation
            LIMIT 10
        """
        results = self._query(cypher, {"ptype": product_type})
        risks = []
        for rec in results:
            risk = {
                "module": rec.get("module"),
                "failure_mode": rec.get("failure_mode"),
                "cause": rec.get("cause"),
                "severity": rec.get("severity"),
                "occurrence": rec.get("occurrence"),
                "detection": rec.get("detection"),
                "mitigation": rec.get("mitigation", "（来自Neo4j）"),
                "source": "Neo4j"
            }
            if all(k in risk for k in ["severity","occurrence","detection"]):
                risk["RPN"] = risk["severity"] * risk["occurrence"] * risk["detection"]
            risks.append(risk)
        return sorted(risks, key=lambda x: x.get("RPN", 0), reverse=True)[:10]

    def get_product_decomposition(self, product_name: str, description: str) -> Dict:
        return {}

    def get_mitigation(self, module: str, failure_mode: str) -> str:
        if not self.driver: return ""
        cypher = """
            MATCH (m:Module {name: $module})-[:HAS_FAILURE]->(f:FailureMode {name: $failure})
            OPTIONAL MATCH (f)-[:MITIGATED_BY]->(mit:Mitigation)
            RETURN mit.text AS mitigation
            LIMIT 1
        """
        results = self._query(cypher, {"module": module, "failure": failure_mode})
        if results and results[0].get("mitigation"):
            return results[0]["mitigation"]
        return ""

    def get_knowledge_by_category(self, category: str) -> List[str]:
        if not self.driver: return []
        lang = st.session_state.lang
        if lang == "zh":
            cypher = "MATCH (k:Knowledge {category: $cat}) RETURN k.content AS content"
        else:
            cypher = "MATCH (k:Knowledge {category: $cat}) RETURN k.content_en AS content"
        results = self._query(cypher, {"cat": category})
        return [r["content"] for r in results if r.get("content")]

    def add_knowledge(self, category: str, content: str):
        if not self.driver: return
        lang = st.session_state.lang
        if lang == "zh":
            zh_text = content
            en_text = translate_text(content, "en")
        else:
            en_text = content
            zh_text = translate_text(content, "zh")
        import uuid
        node_id = str(uuid.uuid4())
        with self.driver.session() as session:
            session.run("CREATE (k:Knowledge {id: $id, category: $cat, content: $zh, content_en: $en})",
                        {"id": node_id, "cat": category, "zh": zh_text, "en": en_text})

    def delete_knowledge(self, category: str, content: str):
        if not self.driver: return
        lang = st.session_state.lang
        if lang == "zh":
            cypher = "MATCH (k:Knowledge {category: $cat, content: $cont}) DELETE k"
        else:
            cypher = "MATCH (k:Knowledge {category: $cat, content_en: $cont}) DELETE k"
        with self.driver.session() as session:
            session.run(cypher, {"cat": category, "cont": content})

    def clear_knowledge_category(self, category: str):
        if not self.driver: return
        with self.driver.session() as session:
            session.run("MATCH (k:Knowledge {category: $cat}) DELETE k", {"cat": category})

    def get_all_knowledge(self) -> Dict[str, List[str]]:
        if not self.driver: return {}
        lang = st.session_state.lang
        if lang == "zh":
            cypher = "MATCH (k:Knowledge) RETURN k.category AS cat, k.content AS cont"
        else:
            cypher = "MATCH (k:Knowledge) RETURN k.category AS cat, k.content_en AS cont"
        results = self._query(cypher)
        knowledge = {}
        for r in results:
            cat = r["cat"]
            knowledge.setdefault(cat, []).append(r["cont"])
        return knowledge

    def search_knowledge(self, keywords: str, limit: int = 5) -> List[str]:
        if not self.driver or not keywords.strip(): return []
        cypher = """
            MATCH (k:Knowledge)
            WHERE k.content CONTAINS $kw OR k.content_en CONTAINS $kw
            RETURN k.content AS zh, k.content_en AS en
            LIMIT $lim
        """
        results = self._query(cypher, {"kw": keywords, "lim": limit})
        lang = st.session_state.lang
        items = []
        for r in results:
            if lang == "zh":
                items.append(r.get("zh", ""))
            else:
                items.append(r.get("en", ""))
        return [item for item in items if item]

    def load_initial_data(self):
        pass

class HybridDatabase(RiskDatabase):
    def __init__(self):
        self.sqlite = SQLiteDatabase()
        self.neo4j = Neo4jDatabase()
        self.neo4j_available = self.neo4j.driver is not None

    def get_risks(self, product_type: str) -> List[Dict]:
        risks_sql = self.sqlite.get_risks(product_type)
        risks_neo = self.neo4j.get_risks(product_type) if self.neo4j_available else []
        seen = set()
        merged = []
        for r in risks_sql + risks_neo:
            key = (r.get("module"), r.get("failure_mode"))
            if key not in seen:
                seen.add(key)
                merged.append(r)
        merged.sort(key=lambda x: x.get("RPN", 0), reverse=True)
        return merged[:10]

    def get_product_decomposition(self, product_name: str, description: str) -> Dict:
        return self.sqlite.get_product_decomposition(product_name, description)

    def get_mitigation(self, module: str, failure_mode: str) -> str:
        sql_mit = self.sqlite.get_mitigation(module, failure_mode)
        if sql_mit and "建议参考" not in sql_mit:
            return sql_mit
        neo_mit = self.neo4j.get_mitigation(module, failure_mode) if self.neo4j_available else ""
        if neo_mit:
            return neo_mit
        return sql_mit

    def get_knowledge_by_category(self, category: str) -> List[str]:
        return self.sqlite.get_knowledge_by_category(category)

    def add_knowledge(self, category: str, content: str):
        self.sqlite.add_knowledge(category, content)
        if self.neo4j_available:
            self.neo4j.add_knowledge(category, content)

    def delete_knowledge(self, category: str, content: str):
        self.sqlite.delete_knowledge(category, content)
        if self.neo4j_available:
            self.neo4j.delete_knowledge(category, content)

    def clear_knowledge_category(self, category: str):
        self.sqlite.clear_knowledge_category(category)
        if self.neo4j_available:
            self.neo4j.clear_knowledge_category(category)

    def get_all_knowledge(self) -> Dict[str, List[str]]:
        return self.sqlite.get_all_knowledge()

    def load_initial_data(self):
        self.sqlite.load_initial_data()
        if self.neo4j_available:
            all_neo = self.neo4j.get_all_knowledge()
            if not any(all_neo.values()):
                conn = self.sqlite.conn
                cursor = conn.cursor()
                cursor.execute("SELECT category, content, content_en FROM knowledge_base")
                rows = cursor.fetchall()
                for cat, zh, en in rows:
                    if zh and en:
                        import uuid
                        node_id = str(uuid.uuid4())
                        with self.neo4j.driver.session() as session:
                            session.run("CREATE (k:Knowledge {id: $id, category: $cat, content: $zh, content_en: $en})",
                                        {"id": node_id, "cat": cat, "zh": zh, "en": en})

    def search_knowledge(self, keywords: str, limit: int = 5) -> List[str]:
        return self.sqlite.search_knowledge(keywords, limit)

def get_database() -> RiskDatabase:
    return HybridDatabase()

# ================== AI 相关函数 ==================
def get_openai_client():
    api_key = st.session_state.temp_api_key if st.session_state.temp_api_key else st.secrets.get("DEEPSEEK_API_KEY", "")
    base_url = st.session_state.temp_base_url if st.session_state.temp_base_url else st.secrets.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
    if not api_key:
        return None, "未配置 API Key"
    return openai.OpenAI(api_key=api_key, base_url=base_url), None

def call_deepseek(prompt: str, max_tokens=4000) -> str:
    client, error = get_openai_client()
    if error:
        return f"AI 调用失败: {error}"
    try:
        model = st.session_state.temp_model if st.session_state.temp_model else st.secrets.get("DEEPSEEK_MODEL", "deepseek-chat")
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=max_tokens
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"AI 调用失败: {str(e)}"

def translate_text(text: str, target_lang: str) -> str:
    if not text or not text.strip():
        return text
    cache_key = f"{text}_{target_lang}"
    if cache_key in st.session_state.translation_cache:
        return st.session_state.translation_cache[cache_key]
    if target_lang == "zh":
        if re.search(r'[\u4e00-\u9fff]', text):
            return text
    else:
        if not re.search(r'[\u4e00-\u9fff]', text):
            return text
    prompt = f"请将以下文本翻译成{'中文' if target_lang == 'zh' else 'English'}，只输出翻译结果：\n\n{text}"
    translated = call_deepseek(prompt, max_tokens=500)
    st.session_state.translation_cache[cache_key] = translated
    return translated

def web_search(query: str, max_results=3) -> str:
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=max_results))
        if not results:
            return "未找到相关结果。"
        output = []
        for r in results:
            output.append(f"- **{r['title']}: {r['body'][:300]}... [来源]({r['href']})")
        return "\n".join(output)
    except Exception as e:
        return f"搜索失败: {str(e)}"

def clean_ai_response(text: str, lang: str = "zh") -> str:
    if lang == "en":
        patterns = [r'^Okay[,.]?\s*\n', r'^As a senior reliability engineer.*?\n', r'^Based on the above information.*?\n', r'^Here is the risk analysis report.*?\n']
    else:
        patterns = [r'^好的[，,].*?\n', r'^作为一名资深可靠性工程师.*?\n', r'^基于以上提供的信息.*?\n', r'^根据您提供的信息.*?\n', r'^以下是对.*?的风险分析报告.*?\n']
    for pat in patterns:
        text = re.sub(pat, '', text, flags=re.IGNORECASE | re.DOTALL)
    lines = text.split('\n')
    if lang == "en":
        if lines and re.match(r'^Okay[,.]?$', lines[0].strip(), re.IGNORECASE):
            text = '\n'.join(lines[1:])
    else:
        if lines and re.match(r'^好的[，,]?$', lines[0].strip()):
            text = '\n'.join(lines[1:])
    return text.strip()

def clean_markdown_text(text: str) -> str:
    text = re.sub(r'\*\*', '', text)
    text = re.sub(r'<br\s*/?>', '\n', text)
    return text

def markdown_to_docx(md_text: str, doc: Document):
    lines = md_text.split('\n')
    i = 0
    in_table = False
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('# '):
            doc.add_heading(clean_markdown_text(line[2:]), level=1)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(clean_markdown_text(line[3:]), level=2)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(clean_markdown_text(line[4:]), level=3)
            i += 1
            continue
        if line.startswith('|') and not in_table:
            in_table = True
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                header_cells = [clean_markdown_text(cell.strip()) for cell in table_lines[0].split('|')[1:-1]]
                if '---' in table_lines[1]:
                    data_lines = table_lines[2:]
                else:
                    data_lines = table_lines[1:]
                num_cols = len(header_cells)
                if num_cols > 0 and data_lines:
                    table = doc.add_table(rows=1+len(data_lines), cols=num_cols)
                    table.style = 'Table Grid'
                    for col_idx, cell_text in enumerate(header_cells):
                        table.cell(0, col_idx).text = cell_text
                        for paragraph in table.cell(0, col_idx).paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    for row_idx, data_line in enumerate(data_lines):
                        cells = [clean_markdown_text(cell.strip()) for cell in data_line.split('|')[1:-1]]
                        for col_idx, cell_text in enumerate(cells):
                            if col_idx < num_cols:
                                table.cell(row_idx+1, col_idx).text = cell_text
                    doc.add_paragraph()
            in_table = False
            continue
        if line:
            p = doc.add_paragraph(clean_markdown_text(line))
            for run in p.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            doc.add_paragraph()
        i += 1

def generate_word_report(product_name: str, product_desc: str, analyst_name: str, analyst_title: str, report_content: str, lang: str = "zh", add_watermark: bool = False) -> BytesIO:
    current_lang = st.session_state.get("lang", "zh")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    if current_lang == "en":
        title_text = "AI-Enabled DQA Product Design Risk Analysis Report"
        url_label = "Report online address:"
        table_labels = {"product_name": "Product Name", "design_desc": "Design Description", "date": "Report Date", "analyst": "Analyst"}
        placeholder = "Not filled"
    else:
        title_text = "AI赋能DQA-产品设计风险分析报告"
        url_label = "报告在线地址："
        table_labels = {"product_name": "产品名称", "design_desc": "设计描述", "date": "报告日期", "analyst": "分析人"}
        placeholder = "未填写"

    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    url_para = doc.add_paragraph(url_label)
    url_para.add_run("https://ai-app-design-dfmea.streamlit.app/").italic = True
    doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_table.cell(0, 0).text = table_labels["product_name"]
    info_table.cell(0, 1).text = product_name
    info_table.cell(1, 0).text = table_labels["design_desc"]
    info_table.cell(1, 1).text = product_desc
    info_table.cell(2, 0).text = table_labels["date"]
    info_table.cell(2, 1).text = datetime.now().strftime("%Y-%m-%d")
    analyst_str = analyst_name if analyst_name else placeholder
    if analyst_title:
        analyst_str += f" ({analyst_title})"
    info_table.cell(3, 0).text = table_labels["analyst"]
    info_table.cell(3, 1).text = analyst_str
    doc.add_paragraph()

    markdown_to_docx(report_content, doc)

    if add_watermark:
        watermark_text = "Confidential - Sample Report - Contact Techlife2027@gmail.com" if current_lang=="en" else "机密 - 样板报告 - 请联系 Techlife2027@gmail.com"
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = header_para.add_run(watermark_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(150, 150, 150)
        run.font.italic = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

def generate_ai_analysis_content(product_name: str, product_desc: str, enable_web: bool, db: RiskDatabase, lang: str = "zh") -> str:
    search_keywords = f"{product_name} {product_desc}"
    kb_items = db.search_knowledge(search_keywords, limit=10)
    kb_text = "\n".join(kb_items) if kb_items else ("No relevant knowledge found." if lang == "en" else "暂无相关经验知识")
    risks = db.get_risks(product_name)
    internal_text = "\n".join([f"- {r['module']}: {r['failure_mode']} (Cause: {r['cause']})" for r in risks[:5]])
    
    web_context = ""
    if enable_web:
        with st.spinner("Searching online..." if lang == "en" else "正在联网搜索..."):
            web_context = web_search(f"{product_name} failure case", max_results=3)
    
    if lang == "en":
        prompt = f"""
You are a senior reliability engineer. Please conduct a risk analysis for the product based on the information below.

Product Name: {product_name}
Design Description: {product_desc}

=== Internal Knowledge Base ===
{kb_text}

=== Product Risk Database ===
{internal_text if internal_text else "None"}

=== Web Search Results ===
{web_context if web_context else "Not enabled"}

IMPORTANT INSTRUCTIONS:
- Output the risk analysis report directly, without any preamble (e.g., "Okay", "Based on the above information").
- Do NOT add any product information table (such as product name, design description, report date, analyst). Start directly with "### 1. Product Decomposition".
- Use ONLY English. Do not output any Chinese characters.
- The report MUST include exactly the following three sections:
### 1. Product Decomposition
### 2. Top 5 Potential Risks (Table: Module, Failure Mode, Cause, Severity, Occurrence, Detection, RPN)
### 3. Key Risk Mitigation Strategies (for the top 3 risks by RPN)

Note: Do not bold module names in the table, and avoid using ** symbols.
"""
    else:
        prompt = f"""
你是一位资深可靠性工程师。请根据以下信息对产品进行风险分析。

产品名称：{product_name}
设计描述：{product_desc}

=== 企业内部知识库（双向检索结果） ===
{kb_text}

=== 产品风险数据库 ===
{internal_text if internal_text else "暂无"}

=== 联网搜索结果 ===
{web_context if web_context else "未启用"}

请直接输出风险分析报告，不要添加任何开场白（如“好的”、“基于以上信息”等）。报告必须包含：
### 1. 产品分解
### 2. Top 5 潜在风险（表格：模块、失效模式、原因、严重度、发生度、探测度、RPN）
### 3. 关键风险缓解策略（针对RPN最高的3项）

注意：表格中的模块名称不要加粗，不要出现 ** 符号。
"""
    raw = call_deepseek(prompt, max_tokens=4000)
    return clean_ai_response(raw, lang)

# ================== 管理员设置弹窗 ==================
@st.dialog("管理员设置", width="large")
def admin_settings_dialog():
    st.subheader("🔐 管理员验证")
    if not st.session_state.admin_logged_in:
        username = st.text_input("用户名")
        password = st.text_input("密码", type="password")
        if st.button("登录"):
            if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
                st.session_state.admin_logged_in = True
                st.rerun()
            else:
                st.error("用户名或密码错误")
        return

    st.success("管理员已登录")
    st.subheader("🌐 联网搜索配置")
    st.session_state.enable_web_search = st.checkbox("启用联网搜索", value=st.session_state.enable_web_search)
    st.markdown("---")
    st.subheader("🗄️ 数据库状态")
    db = st.session_state.database
    neo_available = hasattr(db, 'neo4j_available') and db.neo4j_available
    st.json({
        "当前模式": "混合数据库 (SQLite + Neo4j)",
        "Neo4j 连接": "✅ 已连接" if neo_available else "⚠️ 未连接（仅使用 SQLite）",
        "联网搜索": "启用" if st.session_state.enable_web_search else "禁用",
        "DeepSeek API": "已配置" if (st.session_state.temp_api_key or st.secrets.get("DEEPSEEK_API_KEY")) else "未配置",
        "双向检索": "✅ 已启用"
    })
    st.markdown("---")
    st.subheader("📚 知识库管理（双语）")
    categories = ["光学", "机械", "材料", "热学", "电气", "控制"]
    selected_cat = st.selectbox("选择分类", categories)
    items = db.get_knowledge_by_category(selected_cat)
    st.write(f"共 {len(items)} 条记录")
    if items:
        with st.container(height=400):
            for idx, item in enumerate(items):
                col1, col2 = st.columns([10,1])
                with col1:
                    display_item = item[:150] + "..." if len(item) > 150 else item
                    st.write(f"{idx+1}. {display_item}")
                with col2:
                    if st.button("❌", key=f"del_{selected_cat}_{idx}"):
                        db.delete_knowledge(selected_cat, item)
                        st.rerun()
    else:
        st.info("暂无条目")
    new_item = st.text_area("添加新经验教训（支持中英文，系统会自动翻译存储双语）", height=100)
    if st.button("添加条目"):
        if new_item.strip():
            db.add_knowledge(selected_cat, new_item.strip())
            st.rerun()
    st.markdown("---")
    st.subheader("📥 导出/导入知识库（Excel）")
    if st.button("下载知识库模板 (Excel)"):
        all_zh = st.session_state.database.sqlite.knowledge_zh
        max_len = max((len(all_zh.get(cat, [])) for cat in categories), default=0)
        export_data = {}
        for cat in categories:
            items = all_zh.get(cat, [])
            export_data[cat] = items + [''] * (max_len - len(items))
        df = pd.DataFrame(export_data)
        df.columns = ["光学 / Optical", "机械 / Mechanical", "材料 / Material", "热学 / Thermal", "电气 / Electrical", "控制 / Control"]
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name="知识库", index=False)
        st.download_button(label="下载 Excel 文件", data=output.getvalue(), file_name=f"knowledge_base_{datetime.now().strftime('%Y%m%d')}.xlsx")
    uploaded = st.file_uploader("上传 Excel 文件（覆盖）", type=["xlsx"])
    if uploaded:
        try:
            df = pd.read_excel(uploaded, sheet_name="知识库")
            column_mapping = {
                "光学 / Optical": "光学", "机械 / Mechanical": "机械", "材料 / Material": "材料",
                "热学 / Thermal": "热学", "电气 / Electrical": "电气", "控制 / Control": "控制",
                "光学": "光学", "机械": "机械", "材料": "材料", "热学": "热学", "电气": "电气", "控制": "控制"
            }
            for cat in categories:
                db.clear_knowledge_category(cat)
            for excel_col, cat in column_mapping.items():
                if excel_col in df.columns:
                    items = df[excel_col].dropna().astype(str).tolist()
                    items = [item.strip() for item in items if item.strip()]
                    for item in items:
                        db.add_knowledge(cat, item)
            st.success(f"知识库已更新！共导入 {sum(len(st.session_state.database.sqlite.knowledge_zh[cat]) for cat in categories)} 条记录。")
            st.rerun()
        except Exception as e:
            st.error(f"导入失败：{e}")
    st.markdown("---")
    
    # ========== Report Key 生成器 ==========
    st.subheader("🔑 Report Key 生成器")
    key_type = st.selectbox("选择授权类型", ["试用版", "一级用户", "二级用户", "三级用户", "四级用户", "自定义"])
    custom_uses = None
    custom_months = None
    if key_type == "自定义":
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            custom_uses = st.number_input("使用次数", min_value=1, step=1, value=100)
        with col_c2:
            custom_months = st.number_input("有效期（月）", min_value=1, step=1, value=12)
    custom_key_input = st.text_input("自定义授权码（可选，留空则自动生成）", placeholder="例如：VIP_2026_001")
    if st.button("生成 Report Key"):
        if key_type == "试用版":
            lic_type = "trial"
        elif key_type == "一级用户":
            lic_type = "level1"
        elif key_type == "二级用户":
            lic_type = "level2"
        elif key_type == "三级用户":
            lic_type = "level3"
        elif key_type == "四级用户":
            lic_type = "level4"
        else:
            lic_type = "custom"
        result = generate_report_key(lic_type, custom_uses, custom_months, custom_key_input)
        if result[0] is None:
            st.error(result[3])
        else:
            new_key, max_uses, expiry_str, type_name = result
            st.success(f"已生成 {type_name} Report Key：")
            st.code(new_key, language="text")
            st.write(f"可使用次数：{max_uses} 次，有效期至：{expiry_str[:10]}")
    
    st.markdown("---")
    st.subheader("📋 已生成的所有 Report Key")
    usage_db = load_usage_data()
    records = []
    for key, data in usage_db.items():
        gen_time = data.get("generated_at")
        if gen_time:
            try:
                gen_dt = datetime.fromisoformat(gen_time)
            except:
                gen_dt = datetime.min
        else:
            gen_dt = datetime.min
        records.append({
            "授权码": key,
            "类型": data.get("type", "unknown"),
            "剩余次数": data["remaining"],
            "总使用次数": data.get("total_uses", 0),
            "有效期至": data["expiry"][:10] if data["expiry"] else "永久",
            "生成时间": gen_dt.strftime("%Y-%m-%d %H:%M:%S") if gen_dt != datetime.min else "未知"
        })
    records.sort(key=lambda x: x["生成时间"], reverse=True)
    show_limit = st.selectbox("显示条数", ["最近10条", "最近20条", "最近50条", "全部"], index=0)
    if show_limit == "最近10条":
        limit = 10
    elif show_limit == "最近20条":
        limit = 20
    elif show_limit == "最近50条":
        limit = 50
    else:
        limit = len(records)
    display_records = records[:limit]
    if display_records:
        df = pd.DataFrame(display_records)
        st.dataframe(df, use_container_width=True)
    else:
        st.info("暂无授权码记录")
    if st.button("📥 导出所有授权码为 Excel"):
        if records:
            df_all = pd.DataFrame(records)
            output = BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df_all.to_excel(writer, sheet_name="授权码列表", index=False)
            excel_data = output.getvalue()
            st.download_button(label="点击下载 Excel 文件", data=excel_data, file_name=f"report_keys_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.warning("暂无数据可导出")
    st.markdown("---")
    st.subheader("⚙️ LLM API 临时配置")
    new_key = st.text_input("DeepSeek API Key", value=st.session_state.temp_api_key, type="password")
    new_url = st.text_input("Base URL", value=st.session_state.temp_base_url)
    new_model = st.text_input("Model", value=st.session_state.temp_model)
    if st.button("应用临时配置"):
        st.session_state.temp_api_key = new_key
        st.session_state.temp_base_url = new_url
        st.session_state.temp_model = new_model
        st.rerun()

# ================== 购买对话框 ==================
@st.dialog("购买+解锁", width="large")
def purchase_dialog():
    lang = st.session_state.lang
    if lang == "zh":
        st.markdown("### 选择套餐")
        st.markdown("""
| 套餐 | 价格 | 次数 | 有效期 |
|------|------|------|--------|
| 单次通行 | 18元 / 3美元 | 3次 | 无限制 |
| 50次套餐 | 180元 / 30美元 | 50次 | 1个月 |
| 1000次套餐 | 1200元 / 200美元 | 1000次 | 12个月 |
""")
    else:
        st.markdown("### Select Plan")
        st.markdown("""
| Plan | Price | Credits | Validity |
|------|-------|---------|----------|
| Single Pass | 18 RMB / $3 | 3 uses | Unlimited |
| 50 Credits | 180 RMB / $30 | 50 uses | 1 month |
| 1000 Credits | 1200 RMB / $200 | 1000 uses | 12 months |
""")
    st.markdown("#### 💳 银行卡/数字钱包支付（Stripe）" if lang=="zh" else "#### 💳 Card / Digital Wallet Payment (Stripe)")
    
    if not stripe.api_key:
        st.error("Stripe 未配置，请联系管理员。" if lang=="zh" else "Stripe not configured. Please contact admin.")
        return
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("🎟️ " + (PLANS["single"]["name_zh"] if lang=="zh" else PLANS["single"]["name_en"]) + f"\n${PLANS['single']['price_usd']}", use_container_width=True):
            create_checkout_session("single")
    with col2:
        if st.button("📦 " + (PLANS["50"]["name_zh"] if lang=="zh" else PLANS["50"]["name_en"]) + f"\n${PLANS['50']['price_usd']}", use_container_width=True):
            create_checkout_session("50")
    with col3:
        if st.button("🚀 " + (PLANS["1000"]["name_zh"] if lang=="zh" else PLANS["1000"]["name_en"]) + f"\n${PLANS['1000']['price_usd']}", use_container_width=True):
            create_checkout_session("1000")
    
    st.markdown("#### 🇨🇳 国内支付（微信 / 支付宝）" if lang=="zh" else "#### 🇨🇳 Domestic Payment (WeChat Pay / Alipay)")
    st.info("支持信用卡、微信支付和支付宝。" if lang=="zh" else "Supports credit cards, WeChat Pay and Alipay.")
    st.markdown("支付成功后会自动跳回本页面，授权码将自动激活。" if lang=="zh" else "You will be redirected back after payment, and the license key will be auto-activated.")

def create_checkout_session(plan_key):
    plan = PLANS[plan_key]
    price_id = plan["price_id"]
    base_url = st.secrets.get("APP_URL", "https://ai-app-design-dfmea.streamlit.app")
    success_url = f"{base_url}?order_success=1&plan={plan_key}"
    cancel_url = f"{base_url}"
    try:
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=["card", "wechat_pay", "alipay"],
            line_items=[{"price": price_id, "quantity": 1}],
            mode="payment",
            payment_method_options={"wechat_pay": {"client": "web"}},
            success_url=success_url,
            cancel_url=cancel_url,
            customer_creation="always",
        )
        st.success("支付链接已生成，请点击下方按钮完成支付" if st.session_state.lang=="zh" else "Payment link generated. Click below to pay.")
        button_html = f'<a href="{checkout_session.url}" target="_blank" style="display: block; background-color: #E60000; color: white; font-weight: bold; font-size: 18px; padding: 12px; border-radius: 8px; text-align: center; text-decoration: none; width: 100%;">前往 Stripe 支付页面</a>'
        st.markdown(button_html, unsafe_allow_html=True)
    except Exception as e:
        st.error(f"创建支付会话失败: {e}" if st.session_state.lang=="zh" else f"Failed to create checkout session: {e}")

# ================== 支付成功回调处理 ==================
def handle_payment_callback():
    params = st.query_params
    if "order_success" in params and "plan" in params:
        plan_key = params["plan"]
        if plan_key in PLANS:
            uses = PLANS[plan_key]["uses"]
            months = PLANS[plan_key]["months"]
            new_key, max_uses, expiry_str, _ = generate_report_key("custom", custom_uses=uses, custom_months=months)
            if new_key:
                st.session_state.current_report_key = new_key
                st.session_state.payment_new_key = new_key
                st.session_state.show_payment_dialog = True
                st.query_params.clear()
                st.rerun()
            else:
                st.error("生成授权码失败，请联系管理员。" if st.session_state.lang=="zh" else "Failed to generate license key. Contact admin.")
                st.query_params.clear()
        else:
            st.error("无效的套餐类型。" if st.session_state.lang=="zh" else "Invalid plan type.")
            st.query_params.clear()

# ================== 支付成功弹窗 ==================
def show_payment_success_dialog():
    if st.session_state.get("show_payment_dialog", False):
        @st.dialog("✅ 支付成功")
        def payment_success_dialog():
            lang = st.session_state.lang
            st.markdown("### 您的授权码已生成" if lang=="zh" else "### Your license key has been generated")
            st.code(st.session_state.payment_new_key, language="text")
            st.caption("请妥善保管此授权码，下次使用时可手动复制并粘贴到左侧输入框。" if lang=="zh" else "Please save this license key. You can copy and paste it into the left sidebar next time.")
            st.info("🔑 请复制上方授权码，然后关闭本窗口，回到您原先生成报告的那个窗口，将授权码粘贴到左侧边栏输入框中即可解锁下载。" if lang=="zh" else "🔑 Please copy the license key above, close this window, return to your original report window, and paste the key into the left sidebar to unlock download.")
            if st.button("确定" if lang=="zh" else "OK"):
                st.session_state.show_payment_dialog = False
                st.session_state.payment_new_key = ""
                st.rerun()
        payment_success_dialog()

# ================== 右上角按钮 ==================
col_left, col_spacer, col_zh, col_en, col_gear = st.columns([5, 2, 1.8, 1.8, 1])
with col_zh:
    if st.button("🇨🇳 中文", key="zh_btn", use_container_width=True):
        st.session_state.lang = "zh"
        st.rerun()
with col_en:
    if st.button("🇬🇧 English", key="en_btn", use_container_width=True):
        st.session_state.lang = "en"
        st.rerun()
with col_gear:
    if st.button("⚙️", key="settings_btn", use_container_width=True):
        admin_settings_dialog()

# ================== 多语言文本 ==================
TEXTS = {
    "zh": {
        "title": "🔍 AI+DQA 产品风险分析系统",
        "sidebar_title": "关于系统",
        "basis_items": ["25+年研发管理经验", "AI大模型数据分析", "知识图谱+图神经网络", "DFSS/六西格玛方法论"],
        "analyst_name_label": "分析人姓名",
        "analyst_name_ph": "请输入姓名",
        "analyst_title_label": "分析人头衔（可选）",
        "analyst_title_ph": "例如：研发总监",
        "api_status": "DeepSeek API 状态",
        "api_configured": "✅ 已配置",
        "api_not_configured": "❌ 未配置",
        "contact_info": "📞 **联系：**  \n✉️ 电邮: Techlife2027@gmail.com",
        "input_title": "📝 产品风险分析",
        "product_name": "产品名称",
        "product_name_ph": "例如：高功率LED天棚灯",
        "product_desc": "设计描述",
        "product_desc_ph": "例如：200W COB光源，主动风扇散热，IP65",
        "analyze_btn": "开始AI深度分析",
        "product_name_missing": "请填写产品名称",
        "generating": "AI 正在分析中，请稍候...",
        "footer": "© 2026 Laurence Ku | AI+DQA 风险分析",
        "db_status": "数据库状态",
        "db_connected": "✅ 混合模式 (SQLite + Neo4j)",
        "license_info": "授权信息",
        "remaining_label": "剩余次数",
        "expiry_label": "有效期至",
        "report_key_label": "授权码 (Report Key)",
        "no_license": "未输入授权码，当前为试用模式（剩余次数：{}）",
        "trial_warning": "⚠️ 您还有 {} 次试用机会，输入授权码可解锁无限使用和下载功能。",
        "purchase_button": "💰 购买授权码",
        "download_btn": "📥 下载 Word 报告",
        "need_license": "⚠️ 请先购买授权码后再下载报告。",
    },
    "en": {
        "title": "🔍 AI+DQA Product Risk Analysis",
        "sidebar_title": "About",
        "basis_items": ["25+ years R&D", "AI big data", "Knowledge Graph+GNN", "DFSS/Six Sigma"],
        "analyst_name_label": "Analyst Name",
        "analyst_name_ph": "Enter name",
        "analyst_title_label": "Title (Optional)",
        "analyst_title_ph": "e.g., R&D Director",
        "api_status": "DeepSeek API Status",
        "api_configured": "✅ Configured",
        "api_not_configured": "❌ Not configured",
        "contact_info": "📞 **Contact:**  \n✉️ Email: Techlife2027@gmail.com",
        "input_title": "📝 Product Risk Analysis",
        "product_name": "Product Name",
        "product_name_ph": "e.g., High Bay LED Light",
        "product_desc": "Design Description",
        "product_desc_ph": "e.g., 200W COB, active fan cooling, IP65",
        "analyze_btn": "Start AI Deep Analysis",
        "product_name_missing": "Please enter product name",
        "generating": "AI is analyzing, please wait...",
        "footer": "© 2026 Laurence Ku | AI+DQA Risk Analysis",
        "db_status": "Database Status",
        "db_connected": "✅ Hybrid Mode (SQLite + Neo4j)",
        "license_info": "License Info",
        "remaining_label": "Remaining uses",
        "expiry_label": "Valid until",
        "report_key_label": "Report Key",
        "no_license": "No Report Key. Trial mode (remaining credits: {})",
        "trial_warning": "⚠️ You have {} trial credits left. Enter a license key to unlock unlimited usage.",
        "purchase_button": "💰 Purchase License",
        "download_btn": "📥 Download Word Report",
        "need_license": "⚠️ Please purchase a license before downloading.",
    }
}

lang = st.session_state.lang
t = TEXTS[lang]
st.title(t["title"])

# 初始化数据库
if "database" not in st.session_state:
    st.session_state.database = get_database()
    st.session_state.database.load_initial_data()

# 处理支付回调
handle_payment_callback()
show_payment_success_dialog()

# ================== 侧边栏 ==================
with st.sidebar:
    st.markdown(f"## {t['sidebar_title']}")
    for item in t["basis_items"]:
        st.markdown(f"- {item}")
    st.markdown("---")
    
    analyst_name_input = st.text_input(t["analyst_name_label"], placeholder=t["analyst_name_ph"], key="analyst_name_input")
    analyst_title_input = st.text_input(t["analyst_title_label"], placeholder=t["analyst_title_ph"], key="analyst_title_input")
    st.session_state.analyst_name = analyst_name_input
    st.session_state.analyst_title = analyst_title_input
    if analyst_name_input:
        st.markdown(f"**{t['analyst_name_label']}: {analyst_name_input}**")
        if analyst_title_input:
            st.markdown(f"_{analyst_title_input}_")
    st.markdown("---")
    
    st.markdown(f"**{t['api_status']}**")
    has_api = bool(st.session_state.temp_api_key or st.secrets.get("DEEPSEEK_API_KEY"))
    if has_api:
        st.success(t["api_configured"])
        current_model = st.session_state.temp_model if st.session_state.temp_model else st.secrets.get("DEEPSEEK_MODEL", "deepseek-chat")
        st.caption(f"当前模型: {current_model}")
    else:
        st.error(t["api_not_configured"])
    st.markdown("---")
    
    st.markdown(f"**{t['db_status']}**")
    st.info(t["db_connected"])
    st.markdown("---")
    
    # 授权码输入区域
    st.markdown(f"### 🔑 {t['report_key_label']}")
    new_report_key = st.text_input("", value=st.session_state.current_report_key, type="password", key="report_key_input", placeholder="输入授权码后按 Enter")
    if new_report_key != st.session_state.current_report_key:
        st.session_state.current_report_key = new_report_key
        if new_report_key:
            valid, remaining, expiry_str, _ = activate_license(new_report_key)
            if valid:
                # 成功提示也使用多语言
                if lang == "zh":
                    st.success(f"授权成功！剩余 {remaining} 次，有效期至 {expiry_str[:10]}")
                else:
                    st.success(f"Success! {remaining} uses left, valid until {expiry_str[:10]}")
                st.rerun()
            else:
                st.error("授权码无效或已过期" if lang=="zh" else "Invalid or expired license key")
                st.session_state.current_report_key = ""
                st.rerun()
        else:
            st.rerun()
    
    remaining_str, expiry_str = get_remaining_info(st.session_state.current_report_key)
    st.markdown(f"**{t['license_info']}**")
    st.write(f"{t['remaining_label']}: {remaining_str}")
    if expiry_str not in ("试用剩余次数", "Trial left"):
        st.write(f"{t['expiry_label']}: {expiry_str}")
    if not is_premium_user(st.session_state.current_report_key):
        st.warning(t["trial_warning"].format(st.session_state.trial_uses_left))
    st.markdown("---")
    
    # 购买按钮
    if st.button(t["purchase_button"], use_container_width=True):
        purchase_dialog()
    st.markdown("---")
    st.markdown(t["contact_info"])

# ================== 主界面 ==================
st.markdown(f"### {t['input_title']}")
product_name = st.text_input(t["product_name"], placeholder=t["product_name_ph"], key="product_name_input")
product_desc = st.text_area(t["product_desc"], placeholder=t["product_desc_ph"], height=100, key="product_desc_input")

col_center = st.columns([1, 2, 1])[1]
with col_center:
    st.markdown('<div class="main-analyze">', unsafe_allow_html=True)
    if st.button(t["analyze_btn"], key="main_analyze_btn", type="primary"):
        if not product_name:
            st.error(t["product_name_missing"])
        else:
            if is_premium_user(st.session_state.current_report_key):
                if not consume_usage(st.session_state.current_report_key):
                    st.error("授权码次数已用完或已过期，请购买新授权码。" if lang=="zh" else "License key exhausted or expired. Please purchase a new one.")
                    st.stop()
            else:
                if st.session_state.trial_uses_left <= 0:
                    st.error("试用次数已用完，请购买授权码。" if lang=="zh" else "Trial credits exhausted. Please purchase a license.")
                    purchase_dialog()
                    st.stop()
                consume_usage("")
            
            db = st.session_state.database
            with st.spinner(t["generating"]):
                report_content = generate_ai_analysis_content(
                    product_name, product_desc,
                    st.session_state.enable_web_search,
                    db,
                    lang=st.session_state.lang
                )
                st.session_state.report_content = report_content
                st.session_state.last_product_name = product_name
                st.session_state.last_product_desc = product_desc
                st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# ================== 显示已生成的报告 ==================
if st.session_state.report_content:
    saved_name = st.session_state.get("analyst_name", "")
    saved_title = st.session_state.get("analyst_title", "")
    if saved_name and saved_name.strip():
        author_line = f"分析人：{saved_name.strip()}" + (f" ({saved_title.strip()})" if saved_title.strip() else "") if lang=="zh" else f"Analyst: {saved_name.strip()}" + (f" ({saved_title.strip()})" if saved_title.strip() else "")
    else:
        author_line = "AI生成的风险分析报告" if lang=="zh" else "AI-generated risk analysis report"
    disclaimer_line = "此报告是基于以上提供的有限信息，结合行业数据库和联网搜索结果生成的初步分析，仅供参考。" if lang=="zh" else "This report is a preliminary analysis based on the limited information provided, for reference only."
    full_report_display = f"{author_line}\n\n{disclaimer_line}\n\n{st.session_state.report_content}"
    
    st.markdown("---")
    is_premium = is_premium_user(st.session_state.current_report_key)
    if not is_premium:
        st.markdown(TRIAL_SECURITY_HTML, unsafe_allow_html=True)
    st.markdown('<div class="report-card">', unsafe_allow_html=True)
    st.markdown("### AI赋能DQA-产品设计风险分析报告" if lang=="zh" else "### AI-Enabled DQA Product Design Risk Analysis Report")
    st.markdown(full_report_display)
    st.markdown('</div>', unsafe_allow_html=True)
    
    # 下载按钮
    col_download = st.columns([1,2,1])[1]
    with col_download:
        if st.button(t["download_btn"], use_container_width=True):
            if not is_premium:
                purchase_dialog()
            else:
                word_bytes = generate_word_report(
                    st.session_state.last_product_name,
                    st.session_state.last_product_desc,
                    saved_name, saved_title,
                    st.session_state.report_content,
                    lang=st.session_state.lang,
                    add_watermark=False
                )
                file_name = f"{st.session_state.last_product_name}_风险分析报告_{datetime.now().strftime('%Y%m%d')}.docx" if lang=="zh" else f"{st.session_state.last_product_name}_Risk_Analysis_Report_{datetime.now().strftime('%Y%m%d')}.docx"
                st.download_button(
                    label="📥 确认下载",
                    data=word_bytes,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="real_download"
                )
    if st.button("← 返回重新填写"):
        st.session_state.report_content = None
        st.session_state.last_product_name = ""
        st.session_state.last_product_desc = ""
        st.rerun()

st.markdown("---")
st.caption(t["footer"])
